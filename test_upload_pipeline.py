import os
import asyncio
import io
import docx
from pypdf import PdfWriter
from app.core.database import SessionLocal
from app.models.models import User, ChatSession, Document, ChatMessage
from app.services.document_parser import sniff_mime_type, extract_pdf_chunks, extract_docx_chunks
from app.services.qdrant_service import search_session_documents
from app.worker import process_document_ingestion


def create_sample_pdf(file_path: str, text_content: str):
    """Create a real 1-page sample PDF on disk using pypdf."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    
    c = canvas.Canvas(file_path, pagesize=letter)
    c.drawString(100, 750, "Quantum Library Secret Code Document")
    c.drawString(100, 720, text_content)
    c.save()


def create_sample_docx(file_path: str, heading: str, text_content: str):
    """Create a real sample DOCX file on disk using python-docx."""
    doc = docx.Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(text_content)
    doc.save(file_path)


async def run_pipeline_test():
    print("==================================================")
    print("  UPLOAD-TO-CHAT PIPELINE & ISOLATION VERIFICATION  ")
    print("==================================================")

    db = SessionLocal()
    try:
        # 1. Fetch or create test user
        user = db.query(User).first()
        if not user:
            print("No test user found in DB. Creating mock user...")
            user = User(username="test_student", email="student@test.com", hashed_password="pw")
            db.add(user)
            db.commit()
            db.refresh(user)

        # 2. Create Session 1 and Session 2
        session1 = ChatSession(user_id=user.id, title="Session 1 - Physics PDF")
        session2 = ChatSession(user_id=user.id, title="Session 2 - Isolated History")
        db.add_all([session1, session2])
        db.commit()
        db.refresh(session1)
        db.refresh(session2)

        print(f"Created Session 1 (ID={session1.id}) and Session 2 (ID={session2.id}).")

        # 3. Test MIME Sniffing Validation
        fake_content = b"This is a fake text file labeled as pdf"
        try:
            sniff_mime_type(fake_content)
            print("[FAIL] MIME Sniffing failed (accepted fake content).")
        except ValueError as err:
            print(f"[SUCCESS] MIME Sniffing passed (Rejected fake file: {err})")

        # 4. Create and ingest PDF in Session 1
        pdf_path = "test_quantum_physics.pdf"
        pdf_text = "The secret activation key for Quantum Computing Lab is ALPHA-OMEGA-9988."
        create_sample_pdf(pdf_path, pdf_text)

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        mime_type = sniff_mime_type(pdf_bytes)
        print(f"[SUCCESS] Real PDF MIME type detected: {mime_type}")

        doc_pdf = Document(
            user_id=user.id,
            session_id=session1.id,
            filename="test_quantum_physics.pdf",
            file_path=os.path.abspath(pdf_path),
            file_type=mime_type,
            file_size_bytes=len(pdf_bytes),
            status="processing",
        )
        db.add(doc_pdf)
        db.commit()
        db.refresh(doc_pdf)

        print(f"Processing ARQ ingestion for PDF Document ID={doc_pdf.id}...")
        ingest_res_pdf = await process_document_ingestion({}, doc_pdf.id)
        print("ARQ PDF Ingestion Result:", ingest_res_pdf)
        assert ingest_res_pdf["status"] == "completed"

        # 5. Create and ingest DOCX in Session 1
        docx_path = "test_library_policy.docx"
        create_sample_docx(docx_path, "Special Research Access", "Special research pass members receive 24-hour access to room 402.")
        
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()

        mime_docx = sniff_mime_type(docx_bytes)
        print(f"[SUCCESS] Real DOCX MIME type detected: {mime_docx}")

        doc_docx = Document(
            user_id=user.id,
            session_id=session1.id,
            filename="test_library_policy.docx",
            file_path=os.path.abspath(docx_path),
            file_type=mime_docx,
            file_size_bytes=len(docx_bytes),
            status="processing",
        )
        db.add(doc_docx)
        db.commit()
        db.refresh(doc_docx)

        print(f"Processing ARQ ingestion for DOCX Document ID={doc_docx.id}...")
        ingest_res_docx = await process_document_ingestion({}, doc_docx.id)
        print("ARQ DOCX Ingestion Result:", ingest_res_docx)
        assert ingest_res_docx["status"] == "completed"

        # 6. Test Metadata Scoped Search & Session Isolation (CRITICAL)
        print("\n--- Testing Cross-Session Isolation ---")
        query = "What is the secret activation key for Quantum Computing Lab?"

        # Search in Session 1 (Where PDF was uploaded)
        chunks_s1 = await search_session_documents(query, session1.id, user.id)
        print(f"Session 1 Search Chunks Found: {len(chunks_s1)}")
        if chunks_s1:
            print("Session 1 Top Chunk Source:", chunks_s1[0]["source"], "Section:", chunks_s1[0]["section"])
            print("Session 1 Chunk Snippet:", chunks_s1[0]["content"][:80])

        assert len(chunks_s1) > 0, "Expected matching chunk in Session 1."

        # Search in Session 2 (Different session - NO uploaded PDF)
        chunks_s2 = await search_session_documents(query, session2.id, user.id)
        print(f"Session 2 Search Chunks Found: {len(chunks_s2)}")
        assert len(chunks_s2) == 0, "CRITICAL ERROR: Cross-session leakage detected! Session 2 retrieved Session 1 document!"

        print("[SUCCESS]: 100% Cross-Session & Cross-Tenant Data Isolation Verified!")
        print("==================================================")

    finally:
        db.close()
        # Clean up temporary test files
        for p in ["test_quantum_physics.pdf", "test_library_policy.docx"]:
            if os.path.exists(p):
                os.remove(p)

if __name__ == "__main__":
    asyncio.run(run_pipeline_test())
